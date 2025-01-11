import constants
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from outputs.base_output import BaseOutput


class DocxOutput(BaseOutput):
    def __init__(self):
        self.document = Document()
        # TODO: можно передавать имя файла при запуске линтера,
        # тогда нужно filename передавать как параметр
        self.filename = "report.docx"

    def output_header(self, text: str) -> None:
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = constants.WORD_TIMES_NEW_ROMAN_FONT
        run.font.size = Pt(constants.WORD_HEADER_SIZE)
        run.font.color.rgb = RGBColor(0, 0, 255)

    def output_info(self, text: str, col: int, length: int) -> None:
        paragraph = self.document.add_paragraph()
        before = text[:col]
        highlight = text[col:col+length]
        after = text[col+length:]

        before_run = paragraph.add_run(before)
        before_run.font.name = constants.WORD_TIMES_NEW_ROMAN_FONT
        before_run.font.size = Pt(constants.WORD_MAIN_TEXT_SIZE)

        highlighted_run = paragraph.add_run(highlight)
        highlighted_run.font.name = constants.WORD_TIMES_NEW_ROMAN_FONT
        highlighted_run.font.size = Pt(constants.WORD_MAIN_TEXT_SIZE)
        highlighted_run.font.color.rgb = RGBColor(255, 0, 0)

        after_run = paragraph.add_run(after)
        after_run.font.name = constants.WORD_TIMES_NEW_ROMAN_FONT
        after_run.font.size = Pt(constants.WORD_MAIN_TEXT_SIZE)

    def output_error(self, text: str) -> None:
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = constants.WORD_TIMES_NEW_ROMAN_FONT
        run.font.size = Pt(constants.WORD_MAIN_TEXT_SIZE)
        run.font.color.rgb = RGBColor(255, 128, 0)

    def output_suggestion(self, text: str) -> None:
        fixes_len = len("Варианты исправления:")
        before = text[:fixes_len]
        suggestion = text[fixes_len:]

        paragraph = self.document.add_paragraph()

        before_run = paragraph.add_run(before)
        before_run.font.name = constants.WORD_TIMES_NEW_ROMAN_FONT
        before_run.font.size = Pt(constants.WORD_MAIN_TEXT_SIZE)

        suggestion_run = paragraph.add_run(suggestion)
        suggestion_run.font.name = constants.WORD_TIMES_NEW_ROMAN_FONT
        suggestion_run.font.size = Pt(constants.WORD_MAIN_TEXT_SIZE)
        suggestion_run.font.color.rgb = RGBColor(0, 102, 0)

    def output_newline(self) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p = paragraph._p
        pPr = p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom_border = OxmlElement("w:bottom")

        bottom_border.set(qn("w:val"), "single")
        bottom_border.set(qn("w:sz"), "6")
        bottom_border.set(qn("w:color"), "000000")

        borders.append(bottom_border)
        pPr.append(borders)

    def save(self) -> None:
        self.document.save(self.filename)
